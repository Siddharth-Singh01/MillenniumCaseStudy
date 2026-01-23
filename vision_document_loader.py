#!/usr/bin/env python
# coding: utf-8

"""
Vision-based Document Loader Module

This module provides document loading functionality using vision Model.
It converts PDF and DOCX files to images and uses vision models to extract text
with high accuracy, especially useful for scanned PDFs and complex layouts.

Author: Siddharth Singh
"""

import base64
import io
import logging
from pathlib import Path
from typing import List, Optional
from PIL import Image

from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage
from langchain_core.documents import Document

from config import get_config, AppConfig

# Configure logging
logger = logging.getLogger(__name__)

# Import PyMuPDF (fitz) - required for PDF processing
try:
    import fitz  # PyMuPDF
except ImportError:
    raise ImportError(
        "PyMuPDF (fitz) is required for PDF processing. "
        "Install with: pip install pymupdf"
    )



class VisionDocumentLoader:
    """
    Vision-based document loader for PDF and DOCX files
    Converts documents to images and uses vision API to extract text
    """
    
    def __init__(self, config: Optional[AppConfig] = None):
        """Initialize vision document loader"""
        self.config = config or get_config()
        
        # Get configuration
        vision_cfg = self.config.get_vision_parser_config()
        llm_cfg = self.config.get_llm_config()
        
        # Initialize vision LLM
        llm_kwargs = {
            "api_key": llm_cfg.api_key,
            "model": vision_cfg.vision_model,
            "temperature": 0.0,  # Deterministic extraction
            "max_tokens": 15000  # Increased for longer documents
        }
        if llm_cfg.base_url:
            llm_kwargs["base_url"] = llm_cfg.base_url
        
        self.llm = ChatOpenAI(**llm_kwargs)
        self.dpi = vision_cfg.vision_dpi
        self.detail = vision_cfg.vision_detail
    
    def load_document(self, file_path: Path) -> List[Document]:
        """
        Load document using OpenAI vision API
        
        Args:
            file_path: Path to document file (PDF or DOCX)
            
        Returns:
            List of Document objects with extracted text and metadata
        """
        file_path = Path(file_path)
        
        # Convert document to images
        if file_path.suffix.lower() == '.pdf':
            images = self._pdf_to_images(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            images = self._docx_to_images(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        if not images:
            raise ValueError(f"No images extracted from {file_path}")
        
        # Parse images withvision
        documents = []
        logger.info(f"Processing {len(images)} page(s) from {file_path.name}")
        
        for page_num, image in enumerate(images):
            try:
                # Extract text from image using vision
                text = self._parse_image_with_vision(image)
                
                if not text or len(text.strip()) < 10:
                    logger.warning(f"Page {page_num + 1} returned minimal or no text")
                
                # Create Document object
                doc = Document(
                    page_content=text,
                    metadata={
                        'source': str(file_path),
                        'file_name': file_path.name,
                        'page': page_num,
                        'total_pages': len(images),
                        'extraction_method': 'vision'
                    }
                )
                documents.append(doc)
            except Exception as e:
                logger.error(f"Error processing page {page_num + 1}: {e}")
                # Create empty document to maintain page count
                documents.append(Document(
                    page_content="",
                    metadata={
                        'source': str(file_path),
                        'file_name': file_path.name,
                        'page': page_num,
                        'total_pages': len(images),
                        'extraction_method': 'vision',
                        'error': str(e)
                    }
                ))
        
        logger.info(f"Successfully processed {len(documents)} document(s) from {file_path.name}")
        return documents
    
    def _pdf_to_images(self, pdf_path: Path, dpi: Optional[int] = None) -> List[Image.Image]:
        """
        Convert PDF pages to PIL Images
        
        Args:
            pdf_path: Path to PDF file
            dpi: DPI for conversion (uses config default if None)
            
        Returns:
            List of PIL Images, one per page
        """
        dpi = dpi or self.dpi
        images = []
        
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                # Convert page to pixmap with specified DPI
                matrix = fitz.Matrix(dpi / 72, dpi / 72)
                pix = page.get_pixmap(matrix=matrix)
                # Convert to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append(img)
            doc.close()
        except Exception as e:
            raise ValueError(f"Failed to convert PDF to images: {e}") from e
        
        return images
    
    def _docx_to_images(self, docx_path: Path) -> List[Image.Image]:
        """
        Convert DOCX to images by rendering content using python-docx and PIL
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            List of PIL Images, one per page
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(docx_path)
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {e}") from e
        
        # Extract content from DOCX
        content_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                content_parts.append("\n".join(table_rows))
        
        if not content_parts:
            raise ValueError("DOCX file appears to be empty or unreadable")
        
        # Combine content
        full_text = "\n\n".join(content_parts)
        
        # Render text as images
        images = self._render_text_to_images(full_text)
        
        return images
    
    def _render_text_to_images(self, text: str) -> List[Image.Image]:
        """
        Render text content as images using PIL
        
        Args:
            text: Text content to render
            
        Returns:
            List of PIL Images, split by page size
        """
        from PIL import ImageDraw, ImageFont  # pylint: disable=import-outside-toplevel
        
        # Page dimensions (A4 size at specified DPI)
        dpi = self.dpi
        page_width = int(8.5 * dpi)  # 8.5 inches
        page_height = int(11 * dpi)  # 11 inches
        margin = int(0.5 * dpi)  # 0.5 inch margin
        
        # Text area
        text_width = page_width - (2 * margin)
        text_height = page_height - (2 * margin)
        
        # Font settings
        try:
            # Try to use a system font
            font_size = int(dpi / 6)  # Scale font with DPI
            try:
                font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", font_size)
            except:
                try:
                    font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
                except:
                    font = ImageFont.load_default()
        except:
            font = ImageFont.load_default()
        
        # Split text into lines that fit the page width
        lines = []
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # Simple word wrapping
            words = para.split()
            current_line = []
            current_width = 0
            
            for word in words:
                # Estimate text width (rough approximation)
                word_width = len(word) * font_size * 0.6
                
                if current_width + word_width > text_width and current_line:
                    lines.append(' '.join(current_line))
                    current_line = [word]
                    current_width = word_width
                else:
                    current_line.append(word)
                    current_width += word_width + (font_size * 0.3)  # Add space width
            
            if current_line:
                lines.append(' '.join(current_line))
        
        # Render pages
        images = []
        line_height = int(font_size * 1.5)
        lines_per_page = text_height // line_height
        
        for page_start in range(0, len(lines), lines_per_page):
            page_lines = lines[page_start:page_start + lines_per_page]
            
            # Create image
            img = Image.new('RGB', (page_width, page_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Draw text
            y_position = margin
            for line in page_lines:
                if y_position + line_height > page_height - margin:
                    break
                draw.text((margin, y_position), line, fill='black', font=font)
                y_position += line_height
            
            images.append(img)
        
        return images if images else [Image.new('RGB', (page_width, page_height), color='white')]
    
    def _image_to_base64(self, image: Image.Image) -> str:
        """
        Convert PIL Image to base64 string
        
        Args:
            image: PIL Image object
            
        Returns:
            Base64 encoded string
        """
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    def _parse_image_with_vision(self, image: Image.Image) -> str:
        """
        Parse image using vision API
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text as markdown string
        """
        # Convert image to base64
        base64_image = self._image_to_base64(image)
        
        # Create vision message
        message = HumanMessage(
            content=[
                {
                    "type": "text",
                    "text": """Extract all text from this document page, preserving structure and layout.

Output the content as markdown with proper headers:
- Use # for main sections (e.g., # Name, # Work Experience, # Education)
- Use ## for subsections
- Preserve bullet points and lists
- Maintain table structure if present
- Include all contact information, dates, and details
- Preserve the reading order and logical flow

Extract everything visible on the page, including headers, text blocks, tables, and any structured information."""
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{base64_image}",
                        "detail": self.detail
                    }
                }
            ]
        )
        
        try:
            # Call vision API
            logger.debug("Invoking vision API for text extraction")
            response = self.llm.invoke([message])
            extracted_text = response.content
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                logger.warning("Vision API returned minimal or no text")
            
            return extracted_text
        except Exception as e:
            logger.error(f"vision parsing failed: {e}")
            raise ValueError(f"vision parsing failed: {e}") from e
